#!/usr/bin/env python3
import sys
import re
from pathlib import Path
from typing import List, Tuple

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

LINE_RE = re.compile(r"size=(\d+)KB\s+rtt_avg\(ns\)=(\d+)\s+srv_proc_avg\(ns\)=(\d+)\s+srv_queue_est\(ns\)=(\d+)")

def parse(path: Path):
    rows = []
    for line in path.read_text().splitlines():
        m = LINE_RE.search(line)
        if m:
            size, rtt, cpu, q = map(int, m.groups())
            rows.append((size, rtt, cpu, q))
    rows.sort(key=lambda r: r[0])
    return rows

def plot(rows_a, label_a, rows_b, label_b, key_idx, ylabel, out_png: Path, title: str):
    xs = [r[0] for r in rows_a]
    ya = [rows_a[i][key_idx] / 1e6 for i in range(len(xs))]
    yb = [rows_b[i][key_idx] / 1e6 for i in range(len(xs))]
    plt.figure(figsize=(7,4))
    plt.plot(xs, ya, 'o-', label=label_a)
    plt.plot(xs, yb, 's-', label=label_b)
    plt.xlabel('Payload size (KB)')
    plt.ylabel(ylabel)
    plt.title(title)
    plt.grid(True, alpha=0.3)
    plt.legend()
    out_png.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(out_png)
    plt.close()

def main():
    if len(sys.argv) != 5:
        print('usage: report_ssl_compare.py <label_a> <file_a> <label_b> <file_b>')
        sys.exit(2)
    label_a, file_a, label_b, file_b = sys.argv[1:5]
    rows_a = parse(Path(file_a))
    rows_b = parse(Path(file_b))
    outdir = Path('report_out')
    outdir.mkdir(exist_ok=True)
    rtt_png = outdir / 'rtt_ssl_compare.png'
    cpu_png = outdir / 'cpu_ssl_compare.png'
    plot(rows_a, label_a, rows_b, label_b, 1, 'RTT (ms)', rtt_png, 'RTT (SSL default vs image)')
    plot(rows_a, label_a, rows_b, label_b, 2, 'Server CPU (ms)', cpu_png, 'Server CPU (SSL default vs image)')

    doc = Document()
    doc.add_heading('SSL Default vs Image Spec 비교', 0)
    doc.add_paragraph('두 프로파일 모두 TLS 사용. 차이는 암호 스위트/키 교환/서명 정책이다. 각 페이로드 크기마다 100회 반복 테스트 후 평균값을 사용했다.')
    doc.add_paragraph(f'A: {label_a} — OpenSSL 기본(버전/스위트/그룹 기본값; 일반적으로 TLS1.2/1.3 허용, 스위트/그룹 자동 선택, RSA 인증서)')
    doc.add_paragraph('B: image — TLS1.3 고정, ciphersuites=TLS_AES_128_GCM_SHA256:TLS_AES_256_GCM_SHA384, groups=X25519:P-256, '
                     'sigalgs=ed25519:ecdsa_secp256r1_sha256:rsa_pss_rsae_sha256, 0-RTT 금지, ECDSA-P256 인증서')
    doc.add_heading('RTT', level=1)
    doc.add_picture(str(rtt_png), width=Inches(6.0))
    doc.add_heading('서버 CPU 시간', level=1)
    doc.add_picture(str(cpu_png), width=Inches(6.0))
    outdoc = outdir / 'ssl_compare_report.docx'
    doc.save(outdoc)
    print(f'Wrote {outdoc}')

if __name__ == '__main__':
    main()


