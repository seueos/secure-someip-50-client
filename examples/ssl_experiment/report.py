#!/usr/bin/env python3
import re
import json
from pathlib import Path
from dataclasses import dataclass
from typing import List, Dict

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches


LINE_RE = re.compile(r"size=(\d+)KB\s+rtt_avg\(ns\)=(\d+)\s+srv_proc_avg\(ns\)=(\d+)\s+srv_queue_est\(ns\)=(\d+)")


@dataclass
class Entry:
    size_kb: int
    rtt_ns: int
    cpu_ns: int
    queue_ns: int


def parse_file(path: Path) -> List[Entry]:
    entries: List[Entry] = []
    for line in path.read_text().splitlines():
        m = LINE_RE.search(line)
        if m:
            size, rtt, cpu, q = map(int, m.groups())
            entries.append(Entry(size, rtt, cpu, q))
    entries.sort(key=lambda e: e.size_kb)
    return entries


def plot(entries_a: List[Entry], entries_b: List[Entry], out_png: Path, title: str, key: str):
    xs = [e.size_kb for e in entries_a]
    ya = [getattr(e, key) / 1e6 for e in entries_a]  # ns -> ms
    yb = [getattr(e, key) / 1e6 for e in entries_b]
    plt.figure(figsize=(7, 4))
    plt.plot(xs, ya, 'o-', label='no_ssl')
    plt.plot(xs, yb, 's-', label='image_tls')
    plt.xlabel('Payload size (KB)')
    plt.ylabel({'rtt_ns':'RTT (ms)','cpu_ns':'Server CPU (ms)','queue_ns':'Server queue est (ms)'}[key])
    plt.title(title)
    plt.grid(True, alpha=0.3)
    plt.legend()
    out_png.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(out_png)
    plt.close()


def main():
    no_ssl = Path('/tmp/no_ssl_client.txt')
    img_tls = Path('/tmp/image_tls_client.txt')
    if not no_ssl.exists() or not img_tls.exists():
        raise SystemExit('Missing input files in /tmp. Run run_bench.sh first.')
    a = parse_file(no_ssl)
    b = parse_file(img_tls)
    outdir = Path('report_out')
    rtt_png = outdir / 'rtt.png'
    cpu_png = outdir / 'cpu.png'
    plot(a, b, rtt_png, 'RTT vs Payload size', 'rtt_ns')
    plot(a, b, cpu_png, 'Server CPU time vs Payload size', 'cpu_ns')

    # Prepare quick summary table
    table_data: List[Dict] = []
    for ea, eb in zip(a, b):
        table_data.append({
            'size_kb': ea.size_kb,
            'rtt_no_ssl_ms': round(ea.rtt_ns/1e6, 3),
            'rtt_img_ms': round(eb.rtt_ns/1e6, 3),
            'cpu_no_ssl_ms': round(ea.cpu_ns/1e6, 3),
            'cpu_img_ms': round(eb.cpu_ns/1e6, 3),
        })
    (outdir / 'summary.json').write_text(json.dumps(table_data, indent=2))

    # DOCX report
    doc = Document()
    doc.add_heading('vSomeIP ssl_experiment: RTT & CPU 비교', 0)
    doc.add_paragraph('프로파일: no_ssl vs image_tls (TLS1.3, AES-GCM, X25519/P-256, 0-RTT 금지). 각 페이로드 사이즈별 100회 평균.')
    doc.add_heading('RTT', level=1)
    doc.add_picture(str(rtt_png), width=Inches(6.0))
    doc.add_heading('서버 CPU 시간', level=1)
    doc.add_picture(str(cpu_png), width=Inches(6.0))
    doc.add_heading('요약 테이블 (단위: ms)', level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = 'size_kb'
    hdr[1].text = 'rtt_no_ssl'
    hdr[2].text = 'rtt_image_tls'
    hdr[3].text = 'cpu_no_ssl'
    hdr[4].text = 'cpu_image_tls'
    for row in table_data:
        cells = table.add_row().cells
        cells[0].text = str(row['size_kb'])
        cells[1].text = f"{row['rtt_no_ssl_ms']:.3f}"
        cells[2].text = f"{row['rtt_img_ms']:.3f}"
        cells[3].text = f"{row['cpu_no_ssl_ms']:.3f}"
        cells[4].text = f"{row['cpu_img_ms']:.3f}"
    outdoc = outdir / 'ssl_experiment_report.docx'
    doc.save(outdoc)
    print(f"Report written to {outdoc}")


if __name__ == '__main__':
    main()


